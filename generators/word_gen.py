from docx import Document
from docx.shared import Pt
import io

def crear_docx_carta(nombre_usuario, titulo_oferta, empresa, contenido_carta):
    doc = Document()
    
    # Encabezado
    doc.add_heading('CARTA DE PRESENTACIÓN', 0)
    
    # Datos básicos
    p = doc.add_paragraph()
    p.add_run(f"Candidato: ").bold = True
    p.add_run(f"{nombre_usuario}\n")
    p.add_run(f"Empresa: ").bold = True
    p.add_run(f"{empresa}\n")
    p.add_run(f"Puesto: ").bold = True
    p.add_run(f"{titulo_oferta}\n")
    
    doc.add_paragraph("_" * 50) # Línea divisoria
    
    # Cuerpo de la carta
    parrafo = doc.add_paragraph(contenido_carta)
    
    # Formato de fuente
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    
    # Guardar en memoria para descarga inmediata
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer